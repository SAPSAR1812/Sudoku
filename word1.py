import docx
import pandas as pd

d1=docx.Document('answers.docx')
n=3 
ctr=0
name=""
sid=""
A=pd.DataFrame(columns=['Name','ID','Points'])
m=int(input('Number of submissions'))
for q in range(m):
    s='sub'+str(q)+'.docx'
    d=docx.Document(s)
    for para in d.paragraphs:
        if(para.text[0:5]=="Name:"):
            name=para.text[5:]
        elif(para.text[0:12]=="Shaastra ID:"):
            sid=para.text[12:]
    for i in range(n):
        
        for j in range(0,9):
            for k in range(0,9):
                e=d.tables[i].cell(j,k).text
                e1=d1.tables[i].cell(j,k).text
                if(e==e1):
                    ctr=ctr+1
    A.loc[i]=[name]+[sid]+[ctr-83]
    ctr=0
print(A)
    
    
    
    

                
            
